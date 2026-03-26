"""
CV Parser — extracts text from PDF and DOCX files.
"""

import os
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file using pdfplumber."""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file using python-docx."""
    doc = Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def parse_cv(file_path: str) -> str:
    """
    Auto-detect file format and extract text.
    Supports .pdf and .docx files.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Please upload a PDF or DOCX file.")
