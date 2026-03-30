"""
CV Generator — creates professional DOCX resumes in multiple templates.
"""

import os
import re
from typing import Optional
from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from config import TEMP_DIR


# ── Helpers ──────────────────────────────────────────────────────────────────

def _parse_resume_sections(resume_text: str) -> dict:
    """
    Parse raw resume text into sections.
    Returns dict like {"Contact": "...", "Summary": "...", "Experience": "...", ...}
    """
    sections = {}
    alias_map = {
        "CONTACT": "Contact",
        "SUMMARY": "Professional Summary",
        "OBJECTIVE": "Objective",
        "PROFESSIONAL SUMMARY": "Professional Summary",
        "PROFESSIONAL PROFILE": "Professional Summary",
        "PROFILE": "Professional Summary",
        "EXPERIENCE": "Experience",
        "WORK EXPERIENCE": "Experience",
        "PROFESSIONAL EXPERIENCE": "Experience",
        "EMPLOYMENT HISTORY": "Experience",
        "WORK HISTORY": "Experience",
        "EDUCATION": "Education",
        "SKILLS": "Skills",
        "KEY SKILLS": "Skills",
        "TECHNICAL SKILLS": "Technical Skills",
        "CORE SKILLS": "Technical Skills",
        "TOOLS & TECHNOLOGIES": "Technical Skills",
        "TOOLS AND TECHNOLOGIES": "Technical Skills",
        "TECHNICAL PROFICIENCY": "Technical Skills",
        "PROJECTS": "Projects",
        "CERTIFICATIONS": "Certifications",
        "CERTIFICATES": "Certifications",
        "AWARDS": "Awards",
        "ACHIEVEMENTS": "Achievements",
        "LANGUAGES": "Languages",
        "INTERESTS": "Interests",
        "HOBBIES": "Hobbies",
        "REFERENCES": "References",
        "KEYWORDS": "Keywords",
        "RELEVANT KEYWORDS": "Keywords",
    }
    headers = list(alias_map.keys())
    pattern = r"(?i)^(" + "|".join(re.escape(h) for h in headers) + r")\s*:?\s*$"

    current_section = "Header"
    current_lines = []

    for line in resume_text.split("\n"):
        stripped = line.strip()
        match = re.match(pattern, stripped)
        if match:
            # Save previous section
            if current_lines:
                body = "\n".join(current_lines).strip()
                if body:
                    existing = sections.get(current_section, "")
                    sections[current_section] = "\n\n".join(
                        part for part in [existing, body] if part
                    ).strip()
            current_section = alias_map.get(match.group(1).upper(), match.group(1).title())
            current_lines = []
        else:
            current_lines.append(line)

    # Save last section
    if current_lines:
        body = "\n".join(current_lines).strip()
        if body:
            existing = sections.get(current_section, "")
            sections[current_section] = "\n\n".join(
                part for part in [existing, body] if part
            ).strip()

    return sections


def _add_heading(doc: DocxDocument, text: str, level: int = 1,
                 color: Optional[RGBColor] = None, font_size: Optional[int] = None):
    """Add a styled heading to the document."""
    heading = doc.add_heading(text, level=level)
    if color or font_size:
        for run in heading.runs:
            if color:
                run.font.color.rgb = color
            if font_size:
                run.font.size = Pt(font_size)
    return heading


def _add_body_text(doc: DocxDocument, text: str, font_size: int = 10,
                   bold: bool = False, italic: bool = False):
    """Add a paragraph of body text."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(0)
    return para


def _add_separator(doc: DocxDocument):
    """Add a thin horizontal line separator."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run("─" * 70)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(180, 180, 180)


def _strip_bullet_prefix(line: str) -> tuple[bool, str]:
    """Return whether a line is a bullet and the cleaned text."""
    cleaned = re.sub(r"^[\s\-\*\u2022\u25CF\u25AA\u25A0]+\s*", "", line).strip()
    is_bullet = cleaned != line.strip() and bool(cleaned)
    return is_bullet, cleaned if cleaned else line.strip()


# ── Template 1: One-Page ATS ────────────────────────────────────────────────

def generate_one_page_ats(resume_text: str, filename: str = "cv_1page_ats.docx") -> str:
    """
    Clean single-page ATS-friendly template.
    Narrow margins, compact spacing, keyword-dense.
    """
    doc = Document()

    # Narrow margins for single page
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    sections = _parse_resume_sections(resume_text)

    # Header / Name
    header_text = sections.get("Header", "")
    lines = [l.strip() for l in header_text.split("\n") if l.strip()]
    if lines:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(lines[0])
        run.font.size = Pt(16)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        # Contact info
        if len(lines) > 1:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(" | ".join(lines[1:]))
            contact_run.font.size = Pt(8)
            contact_run.font.color.rgb = RGBColor(100, 100, 100)

    _add_separator(doc)

    # Body sections
    section_order = ["Summary", "Professional Summary", "Objective",
                     "Skills", "Technical Skills",
                     "Experience", "Work Experience", "Professional Experience",
                     "Education", "Projects", "Certifications", "Certificates",
                     "Keywords"]

    for sec_name in section_order:
        if sec_name in sections:
            _add_heading(doc, sec_name.upper(), level=2,
                        color=RGBColor(0, 51, 102), font_size=11)
            for line in sections[sec_name].split("\n"):
                if line.strip():
                    if line.strip().startswith(("•", "-", "●", "▪", "*")):
                        para = doc.add_paragraph(line.strip()[1:].strip(), style="List Bullet")
                        for run in para.runs:
                            run.font.size = Pt(9)
                    else:
                        _add_body_text(doc, line.strip(), font_size=9)

    filepath = os.path.join(TEMP_DIR, filename)
    doc.save(filepath)
    return filepath


# ── Template 2: Two-Page Detailed ────────────────────────────────────────────

def generate_two_page_detailed(resume_text: str, filename: str = "cv_2page_detailed.docx") -> str:
    """
    Two-page professional layout with generous spacing and detail.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    sections = _parse_resume_sections(resume_text)

    # Header
    header_text = sections.get("Header", "")
    lines = [l.strip() for l in header_text.split("\n") if l.strip()]
    if lines:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = name_para.add_run(lines[0])
        run.font.size = Pt(22)
        run.bold = True
        run.font.color.rgb = RGBColor(26, 26, 26)

        if len(lines) > 1:
            for contact_line in lines[1:]:
                cp = doc.add_paragraph()
                cr = cp.add_run(contact_line)
                cr.font.size = Pt(10)
                cr.font.color.rgb = RGBColor(80, 80, 80)
                cp.paragraph_format.space_after = Pt(0)

    _add_separator(doc)

    section_order = ["Summary", "Professional Summary", "Objective",
                     "Experience", "Work Experience", "Professional Experience",
                     "Education", "Skills", "Technical Skills",
                     "Projects", "Certifications", "Certificates",
                     "Awards", "Achievements", "Languages", "Keywords"]

    for sec_name in section_order:
        if sec_name in sections:
            _add_heading(doc, sec_name, level=1,
                        color=RGBColor(0, 102, 153), font_size=14)
            _add_separator(doc)
            for line in sections[sec_name].split("\n"):
                if line.strip():
                    if line.strip().startswith(("•", "-", "●", "▪", "*")):
                        para = doc.add_paragraph(line.strip()[1:].strip(), style="List Bullet")
                        for run in para.runs:
                            run.font.size = Pt(10)
                        para.paragraph_format.space_after = Pt(3)
                    else:
                        _add_body_text(doc, line.strip(), font_size=10)

    filepath = os.path.join(TEMP_DIR, filename)
    doc.save(filepath)
    return filepath


# ── Template 3: Modern Left-Sidebar ──────────────────────────────────────────

def generate_modern_sidebar(resume_text: str, filename: str = "cv_modern_sidebar.docx") -> str:
    """
    Two-column layout: left sidebar with skills/contact, right with experience.
    Uses a table to simulate columns.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    sections = _parse_resume_sections(resume_text)

    # Name header (full width)
    header_text = sections.get("Header", "")
    lines = [l.strip() for l in header_text.split("\n") if l.strip()]
    if lines:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(lines[0])
        run.font.size = Pt(20)
        run.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80)

        if len(lines) > 1:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = contact_para.add_run(" • ".join(lines[1:]))
            cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(100, 100, 100)

    # Two-column table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set column widths (left narrow, right wide)
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    left_cell.width = Inches(2.2)
    right_cell.width = Inches(4.5)

    # LEFT SIDEBAR — Skills, Contact, Education, Certifications
    sidebar_sections = ["Skills", "Technical Skills", "Education",
                        "Certifications", "Certificates", "Languages", "Keywords"]
    left_content = ""
    for sec_name in sidebar_sections:
        if sec_name in sections:
            left_content += f"\n{'━' * 20}\n{sec_name.upper()}\n{'━' * 20}\n"
            left_content += sections[sec_name] + "\n"

    left_para = left_cell.paragraphs[0]
    left_para.text = ""
    for line in left_content.split("\n"):
        p = left_cell.add_paragraph()
        r = p.add_run(line)
        if line.startswith("━") or line.isupper():
            r.bold = True
            r.font.color.rgb = RGBColor(44, 62, 80)
            r.font.size = Pt(9)
        else:
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(60, 60, 60)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)

    # RIGHT SIDE — Summary, Experience, Projects
    main_sections = ["Summary", "Professional Summary", "Objective",
                     "Experience", "Work Experience", "Professional Experience",
                     "Projects", "Awards", "Achievements"]

    right_para = right_cell.paragraphs[0]
    right_para.text = ""
    for sec_name in main_sections:
        if sec_name in sections:
            hp = right_cell.add_paragraph()
            hr = hp.add_run(sec_name.upper())
            hr.bold = True
            hr.font.size = Pt(11)
            hr.font.color.rgb = RGBColor(44, 62, 80)
            hp.paragraph_format.space_before = Pt(6)

            for line in sections[sec_name].split("\n"):
                if line.strip():
                    bp = right_cell.add_paragraph()
                    br = bp.add_run(line.strip())
                    br.font.size = Pt(9)
                    bp.paragraph_format.space_after = Pt(2)

    filepath = os.path.join(TEMP_DIR, filename)
    doc.save(filepath)
    return filepath


# ── Template 4: Classic Clean ────────────────────────────────────────────────

def generate_classic_clean(resume_text: str, filename: str = "cv_classic_clean.docx") -> str:
    """
    Traditional top-down ATS-friendly classic format.
    Clean lines, professional, widely accepted.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    sections = _parse_resume_sections(resume_text)

    # Header
    header_text = sections.get("Header", "")
    lines = [l.strip() for l in header_text.split("\n") if l.strip()]
    if lines:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(lines[0])
        run.font.size = Pt(18)
        run.bold = True
        run.font.name = "Georgia"

        if len(lines) > 1:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(" · ".join(lines[1:]))
            cr.font.size = Pt(9)
            cr.font.name = "Georgia"
            cr.font.color.rgb = RGBColor(80, 80, 80)

    # Thick divider
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = div.add_run("═" * 60)
    dr.font.size = Pt(8)
    dr.font.color.rgb = RGBColor(150, 150, 150)

    section_order = ["Summary", "Professional Summary", "Objective",
                     "Experience", "Work Experience", "Professional Experience",
                     "Skills", "Technical Skills",
                     "Education", "Projects",
                     "Certifications", "Certificates", "Awards", "Keywords"]

    for sec_name in section_order:
        if sec_name in sections:
            # Section header
            sh = doc.add_paragraph()
            sr = sh.add_run(sec_name.upper())
            sr.bold = True
            sr.font.size = Pt(11)
            sr.font.name = "Georgia"
            sr.font.color.rgb = RGBColor(51, 51, 51)
            sh.paragraph_format.space_before = Pt(10)
            sh.paragraph_format.space_after = Pt(2)

            # Thin line under heading
            line_para = doc.add_paragraph()
            line_run = line_para.add_run("─" * 60)
            line_run.font.size = Pt(6)
            line_run.font.color.rgb = RGBColor(200, 200, 200)
            line_para.paragraph_format.space_after = Pt(4)

            for line in sections[sec_name].split("\n"):
                if line.strip():
                    if line.strip().startswith(("•", "-", "●", "▪", "*")):
                        para = doc.add_paragraph(line.strip()[1:].strip(), style="List Bullet")
                        for run in para.runs:
                            run.font.size = Pt(10)
                            run.font.name = "Georgia"
                    else:
                        bp = doc.add_paragraph()
                        br = bp.add_run(line.strip())
                        br.font.size = Pt(10)
                        br.font.name = "Georgia"
                        bp.paragraph_format.space_after = Pt(2)

    filepath = os.path.join(TEMP_DIR, filename)
    doc.save(filepath)
    return filepath


# ── Public API ───────────────────────────────────────────────────────────────

TEMPLATES = {
    "1page_ats": ("📄 1-Page ATS", generate_one_page_ats),
    "2page_detailed": ("📋 2-Page Detailed", generate_two_page_detailed),
    "modern_sidebar": ("🎨 Modern Sidebar", generate_modern_sidebar),
    "classic_clean": ("📝 Classic Clean", generate_classic_clean),
}


def generate_cv(resume_text: str, template_key: str,
                version_name: str = "") -> str:
    """
    Generate a CV using the specified template.
    Returns the file path of the generated DOCX.
    """
    if template_key not in TEMPLATES:
        raise ValueError(f"Unknown template: {template_key}")

    label, generator_fn = TEMPLATES[template_key]
    suffix = f"_{version_name}" if version_name else ""
    filename = f"optimized_cv_{template_key}{suffix}.docx"
    return generator_fn(resume_text, filename)
